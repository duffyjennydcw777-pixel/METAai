import docx

doc = docx.Document(r'C:\Users\Gigabyte\Downloads\Phone Link\AI_native_products_Telegram_Sylectus_full_chat.docx')
with open(r'c:\Users\Gigabyte\.gemini\antigravity\scratch\METAai\chat_raw.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        f.write(p.text + '\n')
print("DONE - extracted", len(doc.paragraphs), "paragraphs")
