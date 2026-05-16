import docx

doc = docx.Document(r'C:\Users\Gigabyte\Downloads\Phone Link\ai_product_strategy_session.docx')

with open(r'c:\Dev\METAai\session_text.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        f.write(para.text + '\n')

# Also extract tables if any
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        f.write(' | '.join(cells) + '\n')
    f.write('\n')

print("Done! Written to session_text.txt")
