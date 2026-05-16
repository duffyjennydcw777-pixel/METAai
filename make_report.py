"""Generate Veronika's practice report as DOCX"""
import os
import json
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def make_doc():
    doc = Document()
    for s in doc.sections:
        s.left_margin = Mm(20)
        s.right_margin = Mm(10)
        s.top_margin = Mm(15)
        s.bottom_margin = Mm(15)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Read all parts in order
    parts_dir = os.path.dirname(__file__)
    sections = []
    import glob
    for p in sorted(glob.glob(os.path.join(parts_dir, "report_part*.json"))):
        with open(p, "r", encoding="utf-8") as f:
            sections.extend(json.load(f))
        print(f"  Loaded: {os.path.basename(p)}")

    for item in sections:
        t = item["type"]
        text = item.get("text", "")

        if t == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(16)
            r.font.name = 'Times New Roman'
        elif t == "heading":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(12)
        elif t == "subheading":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(8)
        elif t == "para":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Mm(12.5)
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
        elif t == "bullet":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(10)
            r = p.add_run("• " + text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
        elif t == "numbered":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Mm(12.5)
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
        elif t == "pagebreak":
            doc.add_page_break()
        elif t == "empty":
            doc.add_paragraph()

    out = os.path.join(os.path.dirname(__file__), "Отчет_Цахилова_В.docx")
    doc.save(out)
    print(f"✓ Saved: {out}")
    print(f"  Size: {os.path.getsize(out)/1024:.1f} KB")

if __name__ == "__main__":
    make_doc()
